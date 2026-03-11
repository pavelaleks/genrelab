#!/usr/bin/env python3
"""
Скрипт для конвертации TECHNICAL.md в формат Word (.docx)

Требования:
    pip install python-docx

Использование:
    python convert_to_word.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_formatted_text(paragraph, text):
    """Добавляет текст с форматированием markdown"""
    # Разбиваем текст на части с форматированием
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            # Жирный текст
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Курсив
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Код
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Обычный текст
            paragraph.add_run(part)

def process_table_line(line, doc, table_data):
    """Обрабатывает строку таблицы"""
    if '|' in line:
        cells = [cell.strip() for cell in line.split('|')]
        # Убираем пустые элементы в начале и конце
        cells = [c for c in cells if c]
        if cells and not all(c.startswith('-') for c in cells):
            table_data.append(cells)
    return table_data

def create_table(doc, table_data):
    """Создает таблицу в документе"""
    if not table_data or len(table_data) < 2:
        return
    
    # Определяем количество столбцов
    num_cols = len(table_data[0])
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Заголовок
    header_cells = table.rows[0].cells
    for i, header in enumerate(table_data[0]):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Данные
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data[:num_cols]):
            row_cells[i].text = cell_data

def convert_markdown_to_word(md_file, docx_file):
    """Конвертирует markdown файл в Word документ"""
    doc = Document()
    
    # Настройка стилей документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Читаем markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lines = []
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        original_line = line
        line = line.rstrip()
        
        # Обработка блоков кода
        if line.strip().startswith('```'):
            if in_code_block:
                # Конец блока кода
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_run = code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(0, 0, 128)
                    # Добавляем рамку
                    code_para.paragraph_format.left_indent = Inches(0.5)
                    code_para.paragraph_format.space_before = Pt(6)
                    code_para.paragraph_format.space_after = Pt(6)
                    code_lines = []
                in_code_block = False
            else:
                # Начало блока кода
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Обработка таблиц
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            table_data = process_table_line(line, doc, table_data)
            i += 1
            continue
        
        if in_table and (not line.strip() or not '|' in line):
            # Конец таблицы
            create_table(doc, table_data)
            table_data = []
            in_table = False
            if not line.strip():
                i += 1
                continue
        
        # Заголовки
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.style.font.size = Pt(18)
            i += 1
            continue
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            heading.style.font.size = Pt(16)
            i += 1
            continue
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            heading.style.font.size = Pt(14)
            i += 1
            continue
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        elif line.startswith('##### '):
            heading = doc.add_heading(line[6:].strip(), level=5)
            i += 1
            continue
        elif line.startswith('###### '):
            heading = doc.add_heading(line[7:].strip(), level=6)
            i += 1
            continue
        
        # Горизонтальная линия
        if line.strip() == '---':
            doc.add_paragraph()
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            i += 1
            continue
        
        # Пустая строка
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # Список с маркером
        if re.match(r'^[\s]*[-*]\s', line):
            indent_level = len(line) - len(line.lstrip())
            item_text = re.sub(r'^[\s]*[-*]\s+', '', line).strip()
            p = doc.add_paragraph(item_text, style='List Bullet')
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(indent_level * 0.25)
            i += 1
            continue
        
        # Нумерованный список
        if re.match(r'^\d+\.\s', line.strip()):
            item_text = re.sub(r'^\d+\.\s+', '', line.strip())
            p = doc.add_paragraph(item_text, style='List Number')
            i += 1
            continue
        
        # Обычный текст
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            i += 1
            continue
        
        i += 1
    
    # Обработка последней таблицы, если она есть
    if in_table and table_data:
        create_table(doc, table_data)
    
    # Сохраняем документ
    doc.save(docx_file)
    print(f"✅ Документ успешно создан: {docx_file}")

if __name__ == '__main__':
    try:
        convert_markdown_to_word('TECHNICAL.md', 'TECHNICAL.docx')
    except ImportError:
        print("❌ Ошибка: библиотека python-docx не установлена")
        print("Установите её командой: pip install python-docx")
    except FileNotFoundError:
        print("❌ Ошибка: файл TECHNICAL.md не найден")
    except Exception as e:
        print(f"❌ Ошибка: {e}")
